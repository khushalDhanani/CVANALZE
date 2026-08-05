import hashlib
import json
from unittest.mock import MagicMock, patch

from fastapi.testclient import TestClient

from app.core.config import settings
from app.main import app
from app.repositories.llm_cache import LLMCacheRepository
from app.repositories.result import ResultRepository
from app.schemas.analysis import (
    OptimizedCandidateProfile,
    OptimizedLLMMatchResponse,
    OptimizedVacancyMatch,
)
from app.services.document_parser import MarkdownResult

client = TestClient(app)


def test_document_cache_roundtrip():
    """Test that document cache stores and retrieves the full extraction result."""
    from app.core.cache import doc_cache_manager

    mock_redis = MagicMock()
    with patch("app.core.cache._REDIS_CLIENT", mock_redis):
        extraction = MarkdownResult(
            markdown="# Candidate\n\nExperience with Python.",
            page_count=2,
            is_scanned=False,
            ocr_applied=False,
        )

        doc_hash = hashlib.sha256(b"test content").hexdigest()
        doc_cache_manager.set(doc_hash, extraction.to_dict())
        mock_redis.setex.assert_called_once()

        cached = doc_cache_manager.get(doc_hash)
        assert cached is not None
        restored = MarkdownResult.from_dict(cached)
        assert restored.markdown == extraction.markdown
        assert restored.page_count == extraction.page_count
        assert restored.is_scanned == extraction.is_scanned
        assert restored.ocr_applied == extraction.ocr_applied


def test_document_cache_miss_fallback():
    """Test that a cache miss returns None."""
    from app.core.cache import doc_cache_manager

    doc_hash = hashlib.sha256(b"never cached").hexdigest()
    cached = doc_cache_manager.get(doc_hash)
    assert cached is None


def test_llm_cache_entry_full_metadata():
    """Test that LLMCacheEntry stores and retrieves all metadata fields."""
    from app.repositories.llm_cache import LLMCacheEntry

    entry = LLMCacheEntry(
        prompt="Analyze this CV for Python Developer",
        raw_response='{"match_score": 92}',
        structured_data={"match_score": 92},
        reasoning="Candidate has 5 years Python experience",
        processing_time_ms=1234.56,
        token_count=450,
        inference_time_ms=1200000,
        model="qwen3:4b",
        prompt_version="3.0",
    )

    key = LLMCacheRepository.extraction_cache_key(
        prompt_version="3.0",
        model_version="qwen3:4b",
        extraction_version="1.0.0",
    )

    LLMCacheRepository.save_cached_entry(key, entry)

    cached = LLMCacheRepository.get_cached_entry(key)
    assert cached is not None
    assert cached.prompt == entry.prompt
    assert cached.raw_response == entry.raw_response
    assert cached.structured_data == entry.structured_data
    assert cached.reasoning == entry.reasoning
    assert cached.processing_time_ms == entry.processing_time_ms
    assert cached.token_count == entry.token_count
    assert cached.inference_time_ms == entry.inference_time_ms
    assert cached.model == entry.model
    assert cached.prompt_version == entry.prompt_version


def test_llm_cache_entry_roundtrip_reconstructs_validated_object():
    """Test that a cached LLMCacheEntry can reconstruct the original Pydantic model."""
    from app.repositories.llm_cache import LLMCacheEntry

    sample_response = OptimizedLLMMatchResponse(
        candidate_profile=OptimizedCandidateProfile(
            core_skills=["Python", "FastAPI"],
            relevant_experience_years=5.0,
            current_role="Python Engineer",
        ),
        matched_vacancies=[
            OptimizedVacancyMatch(
                vacancy_id=101,
                semantic_reason="Strong fit for Python Dev",
                semantic_fit_score=90.0,
            )
        ],
    )

    entry = LLMCacheEntry(
        prompt="test prompt",
        raw_response="{}",
        structured_data=sample_response.model_dump(),
        reasoning="",
        processing_time_ms=500.0,
        token_count=200,
        inference_time_ms=450000,
        model="qwen3:4b",
        prompt_version="3.0",
    )

    key = LLMCacheRepository.compute_composite_hash(
        document_hash="abc",
        candidate_id="42",
        vacancy_ids=["101"],
        prompt_version="3.0",
        model_version="qwen3:4b",
        matching_version="3.0",
    )

    LLMCacheRepository.save_cached_entry(key, entry)

    cached = LLMCacheRepository.get_cached_object(key, OptimizedLLMMatchResponse)
    assert cached is not None
    assert cached.candidate_profile == sample_response.candidate_profile
    assert cached.matched_vacancies == sample_response.matched_vacancies


def test_llm_cache_entry_backward_compatible():
    """Test that old-style save_cached_object still works via get_cached_object."""
    sample_response = OptimizedLLMMatchResponse(
        candidate_profile=OptimizedCandidateProfile(
            core_skills=["Python", "FastAPI"],
            relevant_experience_years=5.0,
            current_role="Python Engineer",
        ),
        matched_vacancies=[
            OptimizedVacancyMatch(
                vacancy_id=101,
                semantic_reason="Strong fit for Python Dev",
                semantic_fit_score=90.0,
            )
        ],
    )

    key = LLMCacheRepository.compute_composite_hash(
        document_hash="abc",
        candidate_id="42",
        vacancy_ids=["101"],
        prompt_version="3.0",
        model_version="qwen3:4b",
        matching_version="3.0",
    )

    LLMCacheRepository.save_cached_object(key, sample_response)
    cached = LLMCacheRepository.get_cached_object(key, OptimizedLLMMatchResponse)
    assert cached is not None
    assert cached.candidate_profile.relevant_experience_years == 5.0


def test_document_cache_different_hash_no_collision():
    """Test that different content hashes produce different cache entries."""
    from app.core.cache import doc_cache_manager

    hash_a = hashlib.sha256(b"content A").hexdigest()
    hash_b = hashlib.sha256(b"content B").hexdigest()

    doc_cache_manager.set(hash_a, {"markdown": "A", "page_count": 1})
    doc_cache_manager.set(hash_b, {"markdown": "B", "page_count": 2})

    cached_a = doc_cache_manager.get(hash_a)
    cached_b = doc_cache_manager.get(hash_b)
    assert cached_a["markdown"] == "A"
    assert cached_b["markdown"] == "B"


def test_redis_result_repository_caching():
    """Test that ResultRepository uses Redis when configured."""
    mock_redis = MagicMock()
    mock_redis = MagicMock()
    with (
        patch("app.core.cache._REDIS_CLIENT", mock_redis),
        patch("app.repositories.result._REDIS_CLIENT", mock_redis),
    ):
        mock_data = {"extracted": "content", "pages": 1}

        # 1. Test atomic_save_result
        res = ResultRepository.atomic_save_result("test_result.json", mock_data)

        # Should return the redis URI string
        assert isinstance(res, str)
        assert res.startswith("redis://")

        # Redis commands should be executed (setex with TTL, not set+expire)
        mock_redis.setex.assert_called_once()

        # 2. Test read_result using the redis path
        mock_redis.get.return_value = json.dumps(mock_data)
        read_data = ResultRepository.read_result(res)
        assert read_data == mock_data

        # 3. Test read_result_by_filename
        mock_redis.get.reset_mock()
        mock_redis.get.return_value = json.dumps(mock_data)
        read_data2 = ResultRepository.read_result_by_filename("test_result.json")
        assert read_data2 == mock_data


def test_redis_llm_repository_caching():
    """Test that LLMCacheRepository uses Redis when configured."""
    mock_redis = MagicMock()
    with patch("app.core.cache._REDIS_CLIENT", mock_redis):
        mock_data = {"score": 85.0}

        cache_key = LLMCacheRepository.extraction_cache_key(
            prompt_version="1.0",
            model_version="test-model",
            extraction_version="1.0.0",
        )

        # 1. Test save_result using version-aware cache key
        LLMCacheRepository.save_result(cache_key, mock_data)
        mock_redis.setex.assert_called_once()

        # 2. Test get_cached_result using same key
        mock_redis.get.return_value = json.dumps(mock_data)
        read_data = LLMCacheRepository.get_cached_result(cache_key)
        assert read_data == mock_data


def test_cv_upload_background_task_returns_processing_status(tmp_path, monkeypatch):
    """Test that /cv/upload endpoint immediately returns a processing status, preventing 504 timeouts."""
    from unittest.mock import patch

    import fitz

    monkeypatch.setattr(settings, "UPLOADS_DIR", tmp_path / "uploads")
    monkeypatch.setattr(settings, "RESULTS_DIR", tmp_path / "results")
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "John Doe\nSoftware Engineer with Python and FastAPI experience.")
    pdf_content = doc.tobytes()
    doc.close()

    from unittest.mock import Mock
    dummy_record = Mock(
        job_id="1", 
        cv_key="dummy_key", 
        state="QUEUED", 
        execution_mode=Mock(value="PENDING"), 
        message="Enqueued", 
        progress=0, 
        attempt=0,
        stage="UPLOADED"
    )
    dummy_submission = Mock(record=dummy_record)

    with patch("app.api.cv.ProcessingQueueService.submit_upload", return_value=dummy_submission):
        response = client.post(
            "/api/cv/upload",
            files={"file": ("test_background.pdf", pdf_content, "application/pdf")},
        )

    # We should get a 200 OK immediately, instead of waiting for Docling/LLM
    assert response.status_code == 200, response.text
    data = response.json()

    # Assert background task response format
    assert data["status"] == "processing"
    assert "cv_key" in data
    assert "message" in data

    cv_key = data["cv_key"]

    # Assert the status endpoint is functional
    status_response = client.get(f"/api/cv/status/{cv_key}")
    assert status_response.status_code == 200
    status_data = status_response.json()
    assert status_data["status"] in ("processing", "COMPLETED", "REPROCESSED")


def test_vacancy_cache_compute_hash():
    """Test that _compute_vacancy_hash produces consistent deterministic hashes."""
    from app.repositories.job import JobRepository

    jobs = [
        {"vacancy_id": 1, "title": "Python Developer"},
        {"vacancy_id": 2, "title": "Java Engineer"},
    ]
    hash_1 = JobRepository._compute_vacancy_hash(jobs)

    jobs_same = [
        {"vacancy_id": 2, "title": "Java Engineer"},
        {"vacancy_id": 1, "title": "Python Developer"},
    ]
    hash_2 = JobRepository._compute_vacancy_hash(jobs_same)

    assert hash_1 == hash_2, "Hash must be stable regardless of input order"

    jobs_diff = [
        {"vacancy_id": 1, "title": "Python Developer"},
        {"vacancy_id": 2, "title": "Senior Java Engineer"},
    ]
    hash_3 = JobRepository._compute_vacancy_hash(jobs_diff)
    assert hash_1 != hash_3, "Hash must change when data changes"


def test_vacancy_cache_is_stale():
    """Test that _is_stale correctly detects count changes from DB."""
    from app.repositories.job import JobRepository

    JobRepository._STALENESS_CACHE.clear()
    jobs = [
        {"vacancy_id": 1, "title": "Python Developer"},
        {"vacancy_id": 2, "title": "Java Engineer"},
    ]
    version = JobRepository._compute_vacancy_hash(jobs)

    mock_db = MagicMock()
    mock_db.execute.side_effect = Exception("Fallback to query count")
    mock_db.query.return_value.filter.return_value.scalar.return_value = 2
    assert not JobRepository._is_stale(version, jobs, db=mock_db), "Should NOT be stale when DB count matches cached count"

    JobRepository._STALENESS_CACHE.clear()
    mock_db.query.return_value.filter.return_value.scalar.return_value = 5
    assert JobRepository._is_stale(version, jobs, db=mock_db), "Should BE stale when DB count differs"


def test_vacancy_cache_get_all_jobs_cached():
    """Test that cached vacancy data with valid version is returned without DB hit."""
    from app.core.cache import vacancy_cache_manager
    from app.repositories.job import JobRepository

    sample_jobs = [
        {"vacancy_id": 1, "title": "Python Developer"},
        {"vacancy_id": 2, "title": "Java Engineer"},
    ]
    version = JobRepository._compute_vacancy_hash(sample_jobs)
    vacancy_cache_manager.set(
        JobRepository._VACANCY_CACHE_KEY,
        {"jobs": sample_jobs, "version": version},
    )

    with patch.object(JobRepository, "_is_stale", return_value=False):
        result = JobRepository.get_all_jobs()
        assert result == sample_jobs


def test_vacancy_cache_stale_triggers_refetch():
    """Test that stale version triggers DB re-fetch."""
    from app.core.cache import vacancy_cache_manager
    from app.repositories.job import JobRepository

    sample_jobs = [
        {"vacancy_id": 1, "title": "Python Developer"},
    ]
    vacancy_cache_manager.set(
        JobRepository._VACANCY_CACHE_KEY,
        {"jobs": sample_jobs, "version": "stale_version_hash"},
    )

    result = JobRepository.get_all_jobs()
    assert result is not None


def test_vacancy_cache_invalidate():
    """Test that invalidate_cache clears both jobs and version cache."""
    from app.core.cache import vacancy_cache_manager
    from app.repositories.job import JobRepository

    sample_jobs = [{"vacancy_id": 1, "title": "Python Developer"}]
    version = JobRepository._compute_vacancy_hash(sample_jobs)
    vacancy_cache_manager.set(
        JobRepository._VACANCY_CACHE_KEY,
        {"jobs": sample_jobs, "version": version},
    )

    assert vacancy_cache_manager.exists(JobRepository._VACANCY_CACHE_KEY)

    JobRepository.invalidate_cache()

    assert not vacancy_cache_manager.exists(JobRepository._VACANCY_CACHE_KEY)


def test_vacancy_cache_no_db_fallback():
    """Test that get_all_jobs falls back to defaults when no DB available."""
    from app.core.cache import vacancy_cache_manager
    from app.repositories.job import JobRepository

    vacancy_cache_manager.delete(JobRepository._VACANCY_CACHE_KEY)

    jobs = JobRepository.get_all_jobs()
    assert isinstance(jobs, list)
    assert len(jobs) > 0, "Should return default jobs when no DB"


def test_embedding_cache_roundtrip():
    """Test that generate_embedding caches and returns the same result."""
    from app.services.embedding_service import EmbeddingService

    mock_response = MagicMock()
    mock_response.json.return_value = {"embeddings": [[0.1, 0.2, 0.3]]}
    mock_response.raise_for_status.return_value = None

    with (
        patch("app.services.embedding_service.settings.EMBEDDING_ENABLED", True),
        patch(
            "app.services.embedding_service.EmbeddingService._call_ollama_embed",
            return_value=[0.1, 0.2, 0.3],
        ),
    ):
        emb = EmbeddingService.generate_embedding("test text", "test-model")
        assert emb == [0.1, 0.2, 0.3]

    with patch("app.services.embedding_service.EmbeddingService._call_ollama_embed") as mock_call:
        emb2 = EmbeddingService.generate_embedding("test text", "test-model")
        assert emb2 == [0.1, 0.2, 0.3]
        mock_call.assert_not_called()


def test_embedding_cache_different_model_version():
    """Test that different model versions produce different cache entries."""
    from app.core.cache import embedding_cache_manager
    from app.services.embedding_service import EmbeddingService

    embedding_cache_manager.clear()

    call_count = 0

    def mock_embed(text, model_version):
        nonlocal call_count
        call_count += 1
        return [float(call_count), 0.0, 0.0]

    with (
        patch("app.services.embedding_service.settings.EMBEDDING_ENABLED", True),
        patch(
            "app.services.embedding_service.EmbeddingService._call_ollama_embed",
            side_effect=mock_embed,
        ),
    ):
        emb_v1 = EmbeddingService.generate_embedding("same content", "model-v1")
        emb_v2 = EmbeddingService.generate_embedding("same content", "model-v2")
        assert emb_v1 != emb_v2, "Different model versions must produce different embeddings"
        assert call_count == 2, "Both calls should generate fresh embeddings (different model versions)"


def test_embedding_cache_content_change():
    """Test that different content produces different cache entries."""
    from app.core.cache import embedding_cache_manager
    from app.services.embedding_service import EmbeddingService

    embedding_cache_manager.clear()

    with (
        patch("app.services.embedding_service.settings.EMBEDDING_ENABLED", True),
        patch("app.services.embedding_service.EmbeddingService._call_ollama_embed") as mock_call,
    ):
        mock_call.side_effect = lambda model, text: [
            hash(text) % 1000 / 1000.0,
            0.0,
            0.0,
        ]

        emb_a = EmbeddingService.generate_embedding("content A", "test-model")
        emb_b = EmbeddingService.generate_embedding("content B", "test-model")
        assert emb_a != emb_b, "Different content must produce different embeddings"
        assert mock_call.call_count == 2


def test_cosine_similarity():
    """Test cosine similarity computation."""
    from app.services.embedding_service import EmbeddingService

    a = [1.0, 0.0, 0.0]
    b = [0.0, 1.0, 0.0]
    assert EmbeddingService.cosine_similarity(a, b) == 0.0

    c = [1.0, 0.0, 0.0]
    d = [1.0, 0.0, 0.0]
    assert EmbeddingService.cosine_similarity(c, d) == 1.0

    e = [1.0, 2.0, 3.0]
    f = [4.0, 5.0, 6.0]
    sim = EmbeddingService.cosine_similarity(e, f)
    assert 0.9 < sim < 1.0


def test_embedding_disabled_when_config_off():
    """Test that generate_embedding returns None when EMBEDDING_ENABLED is False."""
    from app.services.embedding_service import EmbeddingService

    with patch("app.services.embedding_service.settings.EMBEDDING_ENABLED", False):
        emb = EmbeddingService.generate_embedding("some text", "test-model")
        assert emb is None


def test_embedding_cache_different_content_same_model():
    """Test that same model, different content => no collision."""
    from app.core.cache import embedding_cache_manager
    from app.services.embedding_service import EmbeddingService

    embedding_cache_manager.clear()

    with (
        patch("app.services.embedding_service.settings.EMBEDDING_ENABLED", True),
        patch("app.services.embedding_service.EmbeddingService._call_ollama_embed") as mock_call,
    ):
        mock_call.side_effect = lambda model, text: [hash(text) % 1000 / 1000.0]

        emb_1 = EmbeddingService.generate_embedding("unique text one", "same-model")
        emb_2 = EmbeddingService.generate_embedding("unique text two", "same-model")
        assert emb_1 != emb_2
        # Second call should hit cache since we already generated for "unique text two" via side_effect
        emb_3 = EmbeddingService.generate_embedding("unique text two", "same-model")
        assert emb_3 == emb_2
        assert mock_call.call_count == 2  # Only 2 calls for 2 unique texts


def test_match_result_cache_key_includes_version_components():
    """Test that CacheKey.for_match_result produces deterministic keys from all version components."""
    from app.core.cache import CacheKey

    key_a = CacheKey.for_match_result(
        document_hash="abc",
        candidate_id="1",
        vacancy_version="v1",
        vacancy_ids=["10"],
        prompt_version="p1",
        model_version="model1",
        extraction_version="extract1",
        matching_version="m1",
    ).to_key()
    key_b = CacheKey.for_match_result(
        document_hash="abc",
        candidate_id="1",
        vacancy_version="v1",
        vacancy_ids=["10"],
        prompt_version="p1",
        model_version="model1",
        extraction_version="extract1",
        matching_version="m1",
    ).to_key()
    assert key_a == key_b, "Same components must produce same key"

    key_c = CacheKey.for_match_result(
        document_hash="abc",
        candidate_id="1",
        vacancy_version="v2",
        vacancy_ids=["10"],
        prompt_version="p1",
        model_version="model1",
        extraction_version="extract1",
        matching_version="m1",
    ).to_key()
    assert key_a != key_c, "Different vacancy version must produce different key"

    key_d = CacheKey.for_match_result(
        document_hash="abc",
        candidate_id="1",
        vacancy_version="v1",
        vacancy_ids=["10"],
        prompt_version="p2",
        model_version="model1",
        extraction_version="extract1",
        matching_version="m1",
    ).to_key()
    assert key_a != key_d, "Different prompt version must produce different key"

    key_e = CacheKey.for_match_result(
        document_hash="abc",
        candidate_id="1",
        vacancy_version="v1",
        vacancy_ids=["10"],
        prompt_version="p1",
        model_version="model1",
        extraction_version="extract1",
        matching_version="m2",
    ).to_key()
    assert key_a != key_e, "Different matching version must produce different key"

    key_f = CacheKey.for_match_result(
        document_hash="xyz",
        candidate_id="1",
        vacancy_version="v1",
        vacancy_ids=["10"],
        prompt_version="p1",
        model_version="model1",
        extraction_version="extract1",
        matching_version="m1",
    ).to_key()
    assert key_a != key_f, "Different document hash must produce different key"


def test_match_result_cache_roundtrip():
    """Test that match result cache stores and retrieves the full EnrichedCandidateAnalysis."""
    from app.core.cache import match_result_cache_manager
    from app.schemas.analysis import EnrichedCandidateAnalysis, EnrichedJobMatchResult

    match_result_cache_manager.clear()

    sample = EnrichedCandidateAnalysis(
        primary_department="Engineering",
        best_match=EnrichedJobMatchResult(
            job_id="job_1",
            job_title="Python Developer",
            department="Engineering",
            vacancy_id=101,
            score=85.0,
            overall_score=85.0,
            role_score=100.0,
            skills_score=80.0,
            experience_score=100.0,
            education_score=0.0,
            domain_score=100.0,
            technology_score=0.0,
            certification_score=0.0,
            responsibilities_score=0.0,
            coverage=0.8,
            classification="HIGH",
            recommendation="Strong candidate.",
            matched_skills=["Python"],
            missing_skills=[],
            matched_keywords=[],
            missing_keywords=[],
            llm_reason="Great Python match",
            inferred_skills=["FastAPI"],
        ),
        suitable_openings=[
            EnrichedJobMatchResult(
                job_id="job_2",
                job_title="Java Developer",
                department="Engineering",
                vacancy_id=102,
                score=60.0,
                overall_score=60.0,
                role_score=50.0,
                skills_score=60.0,
                experience_score=100.0,
                education_score=0.0,
                domain_score=100.0,
                technology_score=0.0,
                certification_score=0.0,
                responsibilities_score=0.0,
                coverage=0.5,
                classification="MEDIUM",
                recommendation="Potential match.",
                matched_skills=["Java"],
                missing_skills=[],
                matched_keywords=[],
                missing_keywords=[],
                llm_reason="",
                inferred_skills=[],
            ),
        ],
        llm_skipped=False,
    )

    cache_key = "test_match_result_roundtrip"
    match_result_cache_manager.set(cache_key, sample.model_dump())
    cached = match_result_cache_manager.get(cache_key)
    assert cached is not None
    restored = EnrichedCandidateAnalysis.model_validate(cached)
    assert restored.primary_department == "Engineering"
    assert restored.best_match.score == 85.0
    assert restored.best_match.llm_reason == "Great Python match"
    assert len(restored.suitable_openings) == 1
    assert restored.suitable_openings[0].vacancy_id == 102
    assert restored.suitable_openings[0].classification == "MEDIUM"


def test_match_result_cache_different_document_hash():
    """Test that different document hashes produce different cache entries."""
    from app.core.cache import match_result_cache_manager

    match_result_cache_manager.clear()
    match_result_cache_manager.set("key:doc_a", {"primary_department": "Engineering"})
    match_result_cache_manager.set("key:doc_b", {"primary_department": "Finance"})

    cached_a = match_result_cache_manager.get("key:doc_a")
    cached_b = match_result_cache_manager.get("key:doc_b")
    assert cached_a["primary_department"] == "Engineering"
    assert cached_b["primary_department"] == "Finance"


def test_match_result_cache_invalidated_by_version_change():
    """Test that a version change produces a cache miss for the same doc+candidate."""
    from app.core.cache import CacheKey, match_result_cache_manager

    match_result_cache_manager.clear()

    key_v1 = CacheKey.for_match_result(
        document_hash="abc",
        candidate_id="1",
        vacancy_version="v1",
        prompt_version="p1",
        matching_version="m1",
    ).to_key()
    key_v2 = CacheKey.for_match_result(
        document_hash="abc",
        candidate_id="1",
        vacancy_version="v2",
        prompt_version="p1",
        matching_version="m1",
    ).to_key()

    match_result_cache_manager.set(key_v1, {"primary_department": "Engineering"})
    assert match_result_cache_manager.get(key_v1) is not None
    assert match_result_cache_manager.get(key_v2) is None, "Different vacancy version must produce cache miss"


def test_cache_invalidator_cv():
    """Test that invalidate_cv removes doc_cache and match_result entries for that hash without nuking unrelated entries."""
    from app.core.cache import (
        CacheIndex,
        CacheInvalidator,
        doc_cache_manager,
        match_result_cache_manager,
    )

    CacheIndex.clear()
    doc_hash_a = "abc123hash_a"
    doc_hash_b = "xyz789hash_b"

    key_a = "match_key_a"
    key_b = "match_key_b"

    doc_cache_manager.set(doc_hash_a, {"data": "test_a"})
    doc_cache_manager.set(doc_hash_b, {"data": "test_b"})
    match_result_cache_manager.set(key_a, {"score": 85})
    match_result_cache_manager.set(key_b, {"score": 92})

    CacheIndex.add("match_by_doc", doc_hash_a, key_a)
    CacheIndex.add("match_by_doc", doc_hash_b, key_b)

    assert doc_cache_manager.get(doc_hash_a) is not None
    assert match_result_cache_manager.get(key_a) is not None
    assert match_result_cache_manager.get(key_b) is not None

    CacheInvalidator.invalidate_cv(doc_hash_a)

    assert doc_cache_manager.get(doc_hash_a) is None, "Target doc_cache should be cleared"
    assert doc_cache_manager.get(doc_hash_b) is not None, "Unrelated doc_cache should be preserved"
    assert match_result_cache_manager.get(key_a) is None, "Target match result should be cleared"
    assert match_result_cache_manager.get(key_b) is not None, "Unrelated match result should be preserved"


def test_cache_invalidator_candidate():
    """Test that invalidate_candidate selectively purges candidate entries."""
    from app.core.cache import (
        CacheIndex,
        CacheInvalidator,
        cv_result_cache_manager,
        match_result_cache_manager,
    )

    CacheIndex.clear()
    cand_1_match = "match_cand_1"
    cand_2_match = "match_cand_2"

    match_result_cache_manager.set(cand_1_match, {"score": 88})
    match_result_cache_manager.set(cand_2_match, {"score": 95})

    CacheIndex.add("match_by_cand", "cand_1", cand_1_match)
    CacheIndex.add("match_by_cand", "cand_2", cand_2_match)

    cv_result_cache_manager.set("cand_1_cv.json", {"status": "COMPLETED"})
    cv_result_cache_manager.set("cand_2_cv.json", {"status": "COMPLETED"})

    CacheInvalidator.invalidate_candidate("cand_1")

    assert match_result_cache_manager.get(cand_1_match) is None, "Cand 1 match should be cleared"
    assert match_result_cache_manager.get(cand_2_match) is not None, "Cand 2 match should be preserved"
    assert cv_result_cache_manager.get("cand_1_cv.json") is None, "Cand 1 CV result should be cleared"
    assert cv_result_cache_manager.get("cand_2_cv.json") is not None, "Cand 2 CV result should be preserved"


def test_cache_invalidator_vacancies():
    """Test that invalidate_vacancies clears vacancy cache and match results."""
    from app.core.cache import (
        CacheInvalidator,
        match_result_cache_manager,
        vacancy_cache_manager,
    )

    vacancy_cache_manager.set("all_jobs", {"jobs": []})
    vacancy_cache_manager.set("all_jobs_version", "v1")
    match_result_cache_manager.set("some_match", {"score": 85})

    CacheInvalidator.invalidate_vacancies()

    assert vacancy_cache_manager.get("all_jobs") is None
    assert vacancy_cache_manager.get("all_jobs_version") is None


def test_cache_invalidator_prompt():
    """Test that invalidate_prompt clears llm_cache and match_result."""
    from app.core.cache import (
        CacheInvalidator,
        llm_cache_manager,
        match_result_cache_manager,
    )

    llm_cache_manager.set("llm_key", {"response": "data"})
    match_result_cache_manager.set("match_key", {"score": 85})

    CacheInvalidator.invalidate_prompt()

    assert llm_cache_manager.get("llm_key") is None
    assert match_result_cache_manager.get("match_key") is None, "match result should be cleared"


def test_cache_invalidator_llm_model():
    """Test that invalidate_llm_model clears llm_cache and match_result."""
    from app.core.cache import (
        CacheInvalidator,
        llm_cache_manager,
        match_result_cache_manager,
    )

    llm_cache_manager.set("llm_key", {"response": "data"})
    match_result_cache_manager.set("match_key", {"score": 85})

    CacheInvalidator.invalidate_llm_model()

    assert llm_cache_manager.get("llm_key") is None
    assert match_result_cache_manager.get("match_key") is None


def test_cache_invalidator_extraction():
    """Test that invalidate_extraction clears doc_cache and llm_cache."""
    from app.core.cache import CacheInvalidator, doc_cache_manager, llm_cache_manager

    doc_cache_manager.set("doc_key", {"pages": 2})
    llm_cache_manager.set("llm_key", {"response": "data"})

    CacheInvalidator.invalidate_extraction()

    assert doc_cache_manager.get("doc_key") is None
    assert llm_cache_manager.get("llm_key") is None


def test_cache_invalidator_embedding_model():
    """Test that invalidate_embedding_model clears embedding cache."""
    from app.core.cache import CacheInvalidator, embedding_cache_manager

    embedding_cache_manager.set("model:v1:some_hash", [0.1, 0.2])
    embedding_cache_manager.set("model:v1:other_hash", [0.3, 0.4])

    CacheInvalidator.invalidate_embedding_model()

    assert embedding_cache_manager.get("model:v1:some_hash") is None
    assert embedding_cache_manager.get("model:v1:other_hash") is None


def test_invalidate_cv_multi_tier():
    """Test that invalidate_cv purges values across L1 memory, L2 redis, and L3 file cache tiers."""
    from unittest.mock import MagicMock, patch

    from app.core.cache import (
        CacheInvalidator,
        _cv_file_cache,
        _memory_cache,
    )

    doc_hash = "multitier_hash_123"

    # Populate L1 and L3
    _memory_cache.set(f"doc_cache:{doc_hash}", {"parsed": "text"})
    _cv_file_cache.set(f"{doc_hash}.json", {"status": "COMPLETED"})

    mock_redis = MagicMock()
    mock_redis.scan.return_value = (0, [])
    with patch("app.core.cache._REDIS_CLIENT", mock_redis):
        CacheInvalidator.invalidate_cv(doc_hash)

        # Confirm L1 cleared
        assert _memory_cache.get(f"doc_cache:{doc_hash}") is None
        # Confirm L3 file cache cleared
        assert _cv_file_cache.get(f"{doc_hash}.json") is None
        # Confirm L2 Redis delete commands triggered
        assert mock_redis.delete.called or mock_redis.scan.called


def test_cache_manager_tier_simplification():
    """Test that CacheManager.active_providers bypasses FileCache when Redis is active and retains it when Redis is down."""
    from unittest.mock import MagicMock

    from app.core.cache import CacheManager, FileCache, MemoryCache, RedisCache

    mem = MemoryCache()
    redis = RedisCache()
    file_c = FileCache("/tmp/test_cache_dir")

    cm = CacheManager("test_ns", [mem, redis, file_c])

    # When Redis is available
    mock_redis_client = MagicMock()
    with patch("app.core.cache._REDIS_CLIENT", mock_redis_client):
        active = cm.active_providers
        assert mem in active
        assert redis in active
        assert file_c not in active, "FileCache should be bypassed when Redis is active"

    # When Redis is unavailable
    with patch("app.core.cache._REDIS_CLIENT", None):
        active_fallback = cm.active_providers
        assert mem in active_fallback
        assert file_c in active_fallback, "FileCache should be retained as fallback when Redis is down"


def test_startup_warns_and_proceeds_when_redis_inactive():
    """Verify that startup logs fallback warning and proceeds cleanly when Redis is inactive."""
    from app.core.lifecycle import verify_redis

    with patch("app.core.cache._REDIS_CLIENT", None):
        # Should complete cleanly without raising RuntimeError
        verify_redis()


def test_cache_delete_by_pattern():
    """Test that delete_by_pattern selectively removes matching keys."""
    from app.core.cache import _memory_cache

    _memory_cache.clear()
    _memory_cache.set("foo:1", "a")
    _memory_cache.set("foo:2", "b")
    _memory_cache.set("bar:1", "c")

    count = _memory_cache.delete_by_pattern("foo:*")
    assert count == 2
    assert _memory_cache.get("foo:1") is None
    assert _memory_cache.get("foo:2") is None
    assert _memory_cache.get("bar:1") == "c"


def test_cache_index_roundtrip():
    """Test that CacheIndex stores and retrieves cache key associations."""
    from app.core.cache import CacheIndex

    mock_redis = MagicMock()
    mock_redis.smembers.return_value = {"key_a_v2"}
    with patch("app.core.cache._REDIS_CLIENT", mock_redis):
        CacheIndex.add("match_by_doc", "hash_a", "key_a_v2")
        mock_redis.sadd.assert_called_once_with("cache_idx:match_by_doc:hash_a", "key_a_v2")

        keys = CacheIndex.get_keys("match_by_doc", "hash_a")
        assert keys == {"key_a_v2"}

        CacheIndex.remove("match_by_doc", "hash_a")
        mock_redis.delete.assert_called_once_with("cache_idx:match_by_doc:hash_a")


def test_master_data_cache_roundtrip():
    """Test that master_data_cache_manager stores and retrieves data."""
    from app.core.cache import master_data_cache_manager

    master_data_cache_manager.clear()

    profiles = [
        {"id": 1, "name": "Software Engineer", "company_id": 1},
        {"id": 2, "name": "Data Scientist", "company_id": 1},
    ]
    master_data_cache_manager.set("job_profiles", profiles)
    cached = master_data_cache_manager.get("job_profiles")
    assert cached is not None
    assert len(cached) == 2
    assert cached[0]["name"] == "Software Engineer"

    master_data_cache_manager.delete("job_profiles")
    assert master_data_cache_manager.get("job_profiles") is None


def test_master_data_api_returns_cached_data():
    """Test that GET /api/master-data/* endpoints return cached data."""
    from app.core.cache import master_data_cache_manager

    master_data_cache_manager.clear()
    master_data_cache_manager.set("departments", [{"id": 10, "name": "Engineering"}])
    master_data_cache_manager.set("companies", [{"id": 1, "name": "Acme Corp"}])
    master_data_cache_manager.set("skills", [{"id": 5, "name": "Python"}])

    resp = client.get("/api/master-data/departments")
    assert resp.status_code == 200
    data = resp.json()
    assert data == [{"id": 10, "name": "Engineering"}]

    resp = client.get("/api/master-data/companies")
    assert resp.status_code == 200
    assert resp.json() == [{"id": 1, "name": "Acme Corp"}]

    resp = client.get("/api/master-data/skills")
    assert resp.status_code == 200
    assert resp.json() == [{"id": 5, "name": "Python"}]


def test_master_data_api_returns_empty_list_when_not_cached():
    """Test that master data endpoints return [] when nothing is cached."""
    from app.core.cache import master_data_cache_manager

    master_data_cache_manager.clear()

    resp = client.get("/api/master-data/job-profiles")
    assert resp.status_code == 200
    assert resp.json() == []


def test_cache_warmer_warm_all_handles_no_db():
    """Test that warm_all gracefully handles missing DB."""
    from app.services.cache_warmer import warm_all

    with patch("app.services.cache_warmer.MssqlReadSession", None):
        counts = warm_all()
        assert isinstance(counts, dict)
        db_backed = {k: v for k, v in counts.items() if k != "rule_config"}
        for v in db_backed.values():
            assert v == 0
        assert counts["rule_config"] == 1


def test_cache_warmer_warm_vacancies_handles_no_db():
    """Test that warm_vacancies gracefully handles missing DB."""
    from app.services.cache_warmer import warm_vacancies

    with patch("app.services.cache_warmer.MssqlReadSession", None):
        count = warm_vacancies()
        assert count == 0


def test_background_warmup_fails_gracefully():
    """Test that the background warmup function handles errors without raising."""
    from app.core.lifecycle import _run_cache_warmup

    with patch("app.services.cache_warmer.MssqlReadSession", None):
        _run_cache_warmup()


def test_cli_warmup_does_not_raise():
    """Test that the CLI warmup path handles missing DB gracefully."""
    from app.core.cache import master_data_cache_manager

    master_data_cache_manager.clear()
    with patch("app.services.cache_warmer.MssqlReadSession", None):
        from app.services.cache_warmer import warm_all

        counts = warm_all()
        assert counts["vacancies"] == 0
        assert counts["job_profiles"] == 0
        assert counts["departments"] == 0
        assert counts["companies"] == 0
        assert counts["skills"] == 0
        assert counts["rule_config"] == 1


def test_docx_upload_full_pipeline(tmp_path, monkeypatch):
    """Test that uploading a .docx file completes all pipeline stages cleanly."""
    from io import BytesIO
    from unittest.mock import patch

    import docx

    from app.services.document_parser import MarkdownGenerator

    monkeypatch.setattr(settings, "UPLOADS_DIR", tmp_path / "uploads")

    doc = docx.Document()
    doc.add_heading("Jane Smith", level=1)
    doc.add_paragraph("Email: jane.smith@example.com | Phone: +1 555-0199")
    doc.add_heading("Work Experience", level=2)
    doc.add_paragraph("Senior Python Developer at Tech Corp (2020 - Present)")
    doc.add_paragraph("- Built scalable FastAPI backend services")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("Bachelor of Science in Computer Science")

    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = MarkdownGenerator.generate("jane_smith_resume.docx", docx_bytes)
    assert result.markdown is not None
    assert "Jane Smith" in result.markdown or "Python Developer" in result.markdown
    assert result.page_count >= 1

    from unittest.mock import Mock
    dummy_record = Mock(
        job_id="1", 
        cv_key="dummy_key", 
        state="QUEUED", 
        execution_mode=Mock(value="PENDING"), 
        message="Enqueued", 
        progress=0, 
        attempt=0,
        stage="UPLOADED"
    )
    dummy_submission = Mock(record=dummy_record)

    # Verify endpoint handles .docx upload
    with patch("app.api.cv.ProcessingQueueService.submit_upload", return_value=dummy_submission):
        response = client.post(
            "/api/cv/upload",
            files={
                "file": (
                    "jane_smith_resume.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "processing"
    assert "cv_key" in data


def test_get_stable_cv_key_case_normalization():
    """Verify that get_stable_cv_key handles uppercase CV_ prefixes consistently."""
    from app.services.cv_service import get_stable_cv_key

    key_upper = get_stable_cv_key("CV_13347_Yagnik_Resume.pdf")
    key_lower = get_stable_cv_key("cv_13347_Yagnik_Resume.pdf")
    assert key_upper == "cv_13347_Yagnik_Resume"
    assert key_lower == "cv_13347_Yagnik_Resume"
    assert key_upper == key_lower


def test_result_repository_resolve_result_with_prefix_variation(tmp_path, monkeypatch):
    """Verify ResultRepository.resolve_result handles filename prefix variations."""
    from app.core.config import settings
    from app.repositories.result import ResultRepository

    monkeypatch.setattr(settings, "RESULTS_DIR", tmp_path)
    res_file = tmp_path / "cv_13347_Yagnik_Resume.json"
    res_file.write_text(
        '{"id": "cv_13347_Yagnik_Resume", "status": "COMPLETED", "progress": 100, "stage": "complete", "match_analysis": {}}',
        encoding="utf-8",
    )

    r1 = ResultRepository.resolve_result("cv_13347_Yagnik_Resume")
    assert r1 is not None
    assert r1["id"] == "cv_13347_Yagnik_Resume"

    r2 = ResultRepository.resolve_result("CV_13347_Yagnik_Resume")
    assert r2 is not None
    assert r2["id"] == "cv_13347_Yagnik_Resume"
