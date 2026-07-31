from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.document_parser import (
    MarkdownGenerator,
    QualityMetricsCalculator,
    ResumeJsonExtractor,
    TextSanitizer,
)


@pytest.fixture
def sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Alex Johnson", level=1)
    doc.add_paragraph("Email: alex.johnson@example.com | Phone: +1-555-0199")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, SQL, Docker, Git")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Senior Backend Engineer at TechCorp (2021 - Present)")
    doc.add_paragraph("Developed high-performance microservices and REST APIs.")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.S. in Computer Science - University of Technology")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_document_parser_extracts_docx(sample_docx_bytes: bytes):
    extraction = MarkdownGenerator.generate("alex_johnson_cv.docx", sample_docx_bytes)

    assert len(extraction.markdown) > 0
    assert "Alex Johnson" in extraction.markdown
    assert "skills" in extraction.markdown.lower()
    assert "Python" in extraction.markdown
    assert extraction.page_count >= 1


def test_text_sanitizer_collapses_spaced_headings_and_cleans_images():
    raw = """
## CONTACT

<!-- image -->

9998209988

## E D U C A T I O N

- BTech Mechanical

## S K I L L S

- Languages: Dart, Python
    """
    clean = TextSanitizer.sanitize(raw)

    assert "<!-- image -->" not in clean
    assert "## EDUCATION" in clean
    assert "## SKILLS" in clean
    assert "BTech Mechanical" in clean


def test_quality_metrics_calculator():
    sample_text = """
## Tarun Gupta
Email: tarun.gupta@example.com
Phone: +91-9998209988
Location: Surat, Gujarat

## PROFILE SUMMARY
Experienced Flutter Developer with 3+ years experience.

## WORK EXPERIENCE
Sr Developer at TechCorp (2022 - Present)
- Developed mobile applications.

## EDUCATION
BTech Computer Science (2018 - 2022)

## SKILLS
Languages: Dart, Python, JavaScript
    """
    metrics = QualityMetricsCalculator.compute(
        text=sample_text,
        page_count=2,
        pdf_type="TEXT_PDF",
        parser_used="docling_fast",
        ocr_applied=False,
    )

    assert metrics["pages"] == 2
    assert metrics["words"] > 20
    assert "contact" in metrics["sections_detected"]
    assert "experience" in metrics["sections_detected"]
    assert "education" in metrics["sections_detected"]
    assert "skills" in metrics["sections_detected"]
    assert metrics["has_email"] is True
    assert metrics["has_phone"] is True
    assert metrics["completeness_score"] >= 0.70


def test_resume_json_extractor():
    sample_text = """
## Tarun Gupta
Email: gtworks05@gmail.com
Phone: 9998209988
Location: Surat, Gujarat

## PROFILE SUMMARY
Results-driven Flutter Developer with 3+ years of experience.

## WORK EXPERIENCE
## Equal SoftTech
Sr Developer (2022 - Present)
- Developed cross-platform apps using Flutter.
- Integrated Firebase and state management.

## EDUCATION
## PANDIT DEENDAYAL ENERGY UNIVERSITY
BTech Mechanical (2018 - 2022)

## SKILLS
Languages: Dart, HTML, CSS, JavaScript, PHP
Frameworks: Flutter, Provider, BLoC
    """
    resume_json = ResumeJsonExtractor.extract(sample_text)

    assert resume_json["contact_info"]["email"] == "gtworks05@gmail.com"
    assert resume_json["contact_info"]["phone"] == "9998209988"
    assert resume_json["summary"].startswith("Results-driven")
    assert len(resume_json["work_experience"]) > 0
    assert len(resume_json["education"]) > 0
    assert "Dart" in resume_json["skills"]["all_skills"]


def test_document_parser_rejects_empty_file():
    with pytest.raises(ValueError, match="0 bytes"):
        MarkdownGenerator.generate("empty.pdf", b"")


def test_document_parser_rejects_invalid_extension():
    with pytest.raises(ValueError, match="Unsupported file extension"):
        MarkdownGenerator.generate("resume.exe", b"sample content")


def test_api_upload_cv_endpoint(sample_docx_bytes: bytes):
    from unittest.mock import patch

    client = TestClient(app)
    with patch("app.api.cv.background_process_cv"):
        response = client.post(
            "/api/cv/upload",
            files={
                "file": (
                    "alex_johnson_cv.docx",
                    sample_docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert response.status_code == 200
        data = response.json()
        assert "cv_key" in data
        assert data["status"] == "processing"
        assert "message" in data


def test_api_upload_rejects_invalid_file_extension():
    client = TestClient(app)
    response = client.post(
        "/api/cv/upload",
        files={"file": ("invalid.exe", b"hello world", "application/octet-stream")},
    )

    assert response.status_code == 400
    assert "Unsupported file extension" in response.json()["detail"]


def test_deterministic_name_extraction_rejects_job_titles():
    sample_text = """
# IT EXECUTIVE
John Doe
Email: john.doe@example.com
Phone: +1-555-0199

## PROFILE SUMMARY
Experienced IT Executive and Systems Administrator.
    """
    resume_json = ResumeJsonExtractor.extract(sample_text)
    contact = resume_json["contact_info"]

    assert contact["name"] == "John Doe"
    assert contact["full_name"] == "John Doe"
    assert contact["name_confidence"] >= 0.85
    assert contact["extraction_source"] in ["header_email_validated", "header_contact_section"]


def test_deterministic_name_extraction_email_fallback():
    sample_text = """
Contact: alex.johnson@company.org
Phone: 9876543210
Location: Chicago, IL

## SUMMARY
Software engineer with 5 years experience.
    """
    resume_json = ResumeJsonExtractor.extract(sample_text)
    contact = resume_json["contact_info"]

    assert contact["name"] == "Alex Johnson"
    assert contact["name_confidence"] == 0.30
    assert contact["extraction_source"] == "email_username_fallback"


def test_deterministic_name_extraction_filename_fallback():
    sample_text = """
Phone: 9876543210
Location: New York, NY

## SUMMARY
Senior Analyst.
    """
    resume_json = ResumeJsonExtractor.extract(sample_text, filename="Jane_Smith_Resume.pdf")
    contact = resume_json["contact_info"]

    assert contact["name"] == "Jane Smith"
    assert contact["name_confidence"] == 0.30
    assert contact["extraction_source"] == "filename_fallback"

